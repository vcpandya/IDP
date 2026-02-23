"""DOCX parser using python-docx for text extraction."""

import logging

from .base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Extract text and metadata from DOCX documents.

    Uses ``python-docx`` to iterate over paragraphs and tables, preserving
    heading hierarchy via markdown-style prefixes (``##``, ``###``, etc.).
    """

    def supported_extensions(self) -> list[str]:
        return ["docx", "doc"]

    def parse(self, file_path: str) -> ParseResult:
        """Parse a DOCX file and return its text content.

        Parameters
        ----------
        file_path:
            Path to the ``.docx`` file.

        Raises
        ------
        ImportError
            If python-docx is not installed.
        """
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            )

        logger.info("Parsing DOCX: %s", file_path)

        document = docx.Document(file_path)
        parts: list[str] = []

        # --- Paragraphs ---
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower() if para.style else ""

            # Map Word heading styles to markdown heading prefixes.
            if style_name.startswith("heading"):
                try:
                    level = int(style_name.replace("heading", "").strip())
                except ValueError:
                    level = 1
                # Heading 1 -> ##, Heading 2 -> ###, etc.
                prefix = "#" * (level + 1)
                parts.append(f"{prefix} {text}")
            elif style_name == "title":
                parts.append(f"# {text}")
            else:
                parts.append(text)

        # --- Tables ---
        for table_idx, table in enumerate(document.tables):
            table_lines = _table_to_markdown(table)
            if table_lines:
                parts.append("")  # blank line before table
                parts.extend(table_lines)
                parts.append("")  # blank line after table

        # --- Metadata ---
        core = document.core_properties
        metadata: dict = {}
        if core.title:
            metadata["title"] = core.title
        if core.author:
            metadata["author"] = core.author
        if core.subject:
            metadata["subject"] = core.subject
        if core.created:
            metadata["creation_date"] = str(core.created)
        if core.modified:
            metadata["modification_date"] = str(core.modified)
        metadata = {k: v for k, v in metadata.items() if v}

        full_text = "\n".join(parts)

        # DOCX does not have a native "page" concept; treat as single page.
        pages = [{"page": 1, "text": full_text}]

        logger.info("DOCX parsing complete: %d characters extracted", len(full_text))

        return ParseResult(
            text=full_text,
            pages=pages,
            metadata=metadata,
            page_count=1,
        )


def _table_to_markdown(table) -> list[str]:
    """Convert a python-docx Table to markdown-formatted lines."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return []

    lines: list[str] = []
    # Header row
    header = rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    # Data rows
    for row in rows[1:]:
        # Pad or trim to match header column count
        padded = row + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(padded[: len(header)]) + " |")

    return lines
