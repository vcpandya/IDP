"""Format batch processing results into DOCX documents."""

import io
import json
import logging
from datetime import datetime, timezone
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

ACCENT_COLOR = RGBColor(0x4F, 0x46, 0xE5)
MUTED_COLOR = RGBColor(0x6B, 0x72, 0x80)
DARK_COLOR = RGBColor(0x1F, 0x29, 0x37)
SEPARATOR_COLOR = RGBColor(0xD1, 0xD5, 0xDB)


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D1D5DB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_title(doc: Document, text: str):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = DARK_COLOR


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT_COLOR if level == 1 else DARK_COLOR


def _add_para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(str(text) if text else "")
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_COLOR
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_label_value(doc: Document, label: str, value):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = DARK_COLOR
    run_val = p.add_run(str(value) if value is not None else "N/A")
    run_val.font.size = Pt(11)
    run_val.font.color.rgb = DARK_COLOR
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_bullet_list(doc: Document, items: list):
    for item in items:
        p = doc.add_paragraph(str(item), style="List Bullet")
        p.paragraph_format.space_after = Pt(3)


def _add_numbered_list(doc: Document, items: list):
    for item in items:
        p = doc.add_paragraph(str(item), style="List Number")
        p.paragraph_format.space_after = Pt(3)


def _add_table(doc: Document, headers: list[str], rows: list[list]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val) if val is not None else ""
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph("")
    return table


def _add_info_box(doc: Document, tool_name: str, filename: Optional[str] = None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]

    tool_display = tool_name.replace("_", " ").title() if tool_name else "Processing"
    lines = [f"Tool: {tool_display}"]
    if filename:
        lines.append(f"Source Document: {filename}")
    lines.append(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")

    for line in lines:
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED_COLOR
        p.paragraph_format.space_after = Pt(1)

    if cell.paragraphs and cell.paragraphs[0].text == "":
        cell._element.remove(cell.paragraphs[0]._element)

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F4F6")
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph("")


def _format_summary(doc: Document, data: dict):
    title = data.get("title") or data.get("document", "Document Summary")
    _add_title(doc, title)

    if data.get("overall_summary"):
        _add_heading(doc, "Summary")
        _add_para(doc, data["overall_summary"])

    sections = data.get("section_summaries", [])
    if sections:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Section Summaries")
        for sec in sections:
            section_title = sec.get("section") or sec.get("title", "Section")
            _add_heading(doc, section_title, level=2)
            _add_para(doc, sec.get("summary") or sec.get("text", ""))

    takeaways = data.get("key_takeaways", [])
    if takeaways:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Key Takeaways")
        _add_bullet_list(doc, takeaways)

    messages = data.get("key_messages", [])
    if messages:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Key Messages")
        _add_bullet_list(doc, messages)


def _format_extract(doc: Document, data: dict):
    _add_title(doc, "Data Extraction Results")

    if data.get("extraction_type"):
        _add_label_value(doc, "Extraction Type", data["extraction_type"])
    if data.get("item_count"):
        _add_label_value(doc, "Items Found", data["item_count"])

    extracted = data.get("data")
    if isinstance(extracted, list):
        _add_horizontal_rule(doc)
        _add_heading(doc, "Extracted Data")
        for i, item in enumerate(extracted, 1):
            if isinstance(item, dict):
                _add_heading(doc, f"Item {i}", level=2)
                for k, v in item.items():
                    if isinstance(v, list):
                        _add_label_value(doc, k.replace("_", " ").title(), ", ".join(str(x) for x in v))
                    elif isinstance(v, dict):
                        _add_label_value(doc, k.replace("_", " ").title(), json.dumps(v, indent=2))
                    else:
                        _add_label_value(doc, k.replace("_", " ").title(), v)
            else:
                _add_para(doc, f"{i}. {item}")
    elif isinstance(extracted, dict):
        _add_horizontal_rule(doc)
        _add_heading(doc, "Extracted Data")
        for k, v in extracted.items():
            if isinstance(v, list):
                _add_heading(doc, k.replace("_", " ").title(), level=2)
                if v and isinstance(v[0], dict):
                    headers = list(v[0].keys())
                    rows = [[item.get(h, "") for h in headers] for item in v if isinstance(item, dict)]
                    _add_table(doc, [h.replace("_", " ").title() for h in headers], rows)
                else:
                    _add_bullet_list(doc, v)
            elif isinstance(v, dict):
                _add_heading(doc, k.replace("_", " ").title(), level=2)
                for sk, sv in v.items():
                    _add_label_value(doc, sk.replace("_", " ").title(), sv)
            else:
                _add_label_value(doc, k.replace("_", " ").title(), v)

    meta = data.get("metadata", {})
    if meta:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Metadata")
        for k, v in meta.items():
            _add_label_value(doc, k.replace("_", " ").title(), v)


def _format_translate(doc: Document, data: dict):
    doc_name = data.get("document", "Translation")
    src = data.get("source_language", "Unknown")
    tgt = data.get("target_language", "Unknown")
    _add_title(doc, f"Translation: {doc_name}")

    info_table = doc.add_table(rows=1, cols=3)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, val) in enumerate([
        ("Source", src), ("Target", tgt),
        ("Sections", str(data.get("total_sections_translated", "N/A")))
    ]):
        cell = info_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.size = Pt(10)
    doc.add_paragraph("")

    sections = data.get("translated_sections", [])
    if sections:
        _add_heading(doc, "Translated Content")
        for sec in sections:
            orig_title = sec.get("original_title")
            trans_title = sec.get("translated_title") or orig_title or "Section"
            _add_heading(doc, trans_title, level=2)
            if orig_title and sec.get("translated_title") and orig_title != sec["translated_title"]:
                p = doc.add_paragraph()
                run = p.add_run(f"Original title: {orig_title}")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = MUTED_COLOR
            _add_para(doc, sec.get("translated_text") or sec.get("text", ""))

    notes = data.get("translation_notes")
    if notes:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Translation Notes")
        if isinstance(notes, list):
            _add_bullet_list(doc, notes)
        else:
            _add_para(doc, notes)


def _format_anonymize(doc: Document, data: dict):
    doc_name = data.get("document", "Anonymization Report")
    _add_title(doc, f"Anonymization Report: {doc_name}")

    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, val) in enumerate([
        ("Style", data.get("redaction_style", "N/A")),
        ("PII Detected", str(data.get("total_pii_detected", 0))),
        ("Redacted", str(data.get("total_redacted", 0))),
    ]):
        cell = summary_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.size = Pt(10)

    cats = data.get("categories_found", [])
    if cats:
        doc.add_paragraph("")
        _add_label_value(doc, "Categories Found", ", ".join(cats) if isinstance(cats, list) else cats)

    pii = data.get("pii_found", [])
    if pii:
        _add_horizontal_rule(doc)
        _add_heading(doc, "PII Detected")
        rows = []
        for item in pii:
            if isinstance(item, dict):
                rows.append([
                    item.get("category", ""),
                    item.get("original", ""),
                    item.get("redacted", ""),
                    str(item.get("confidence", "")),
                    item.get("location", ""),
                ])
        if rows:
            _add_table(doc, ["Category", "Original", "Redacted", "Confidence", "Location"], rows)

    sections = data.get("redacted_sections", [])
    if sections:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Redacted Content")
        for sec in sections:
            if isinstance(sec, dict):
                title = sec.get("section") or sec.get("title", "Section")
                _add_heading(doc, title, level=2)
                _add_para(doc, sec.get("redacted_text") or sec.get("text", ""))
            else:
                _add_para(doc, str(sec))


def _format_audit(doc: Document, data: dict):
    doc_name = data.get("document", "Audit Report")
    _add_title(doc, f"Audit Report: {doc_name}")

    score = data.get("overall_score")
    status = data.get("overall_status", "N/A")
    audit_type = data.get("audit_type", "General")

    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, val) in enumerate([
        ("Type", audit_type),
        ("Score", f"{score}/100" if score is not None else "N/A"),
        ("Status", status),
    ]):
        cell = summary_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.size = Pt(10)
    doc.add_paragraph("")

    if data.get("summary"):
        _add_heading(doc, "Executive Summary")
        _add_para(doc, data["summary"])

    findings = data.get("findings", [])
    if findings:
        _add_horizontal_rule(doc)
        _add_heading(doc, f"Findings ({len(findings)})")
        for i, finding in enumerate(findings, 1):
            if isinstance(finding, dict):
                title = finding.get("title") or finding.get("finding", f"Finding {i}")
                severity = finding.get("severity", "")
                heading_text = f"{i}. {title}"
                if severity:
                    heading_text += f" [{severity.upper()}]"
                _add_heading(doc, heading_text, level=2)
                if finding.get("description"):
                    _add_para(doc, finding["description"])
                if finding.get("impact"):
                    _add_label_value(doc, "Impact", finding["impact"])
                if finding.get("recommendation"):
                    _add_label_value(doc, "Recommendation", finding["recommendation"])
                if finding.get("evidence"):
                    _add_label_value(doc, "Evidence", finding["evidence"])
            else:
                _add_para(doc, f"{i}. {finding}")

    checklist = data.get("checklist_results", [])
    if checklist:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Checklist Results")
        if isinstance(checklist, list) and checklist and isinstance(checklist[0], dict):
            headers = list(checklist[0].keys())
            rows = [[item.get(h, "") for h in headers] for item in checklist]
            _add_table(doc, [h.replace("_", " ").title() for h in headers], rows)
        elif isinstance(checklist, list):
            _add_bullet_list(doc, checklist)

    recs = data.get("recommendations", [])
    if recs:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Recommendations")
        _add_numbered_list(doc, recs)

    stats = data.get("statistics", {})
    if stats and isinstance(stats, dict):
        _add_horizontal_rule(doc)
        _add_heading(doc, "Statistics")
        rows = [[k.replace("_", " ").title(), str(v)] for k, v in stats.items()]
        _add_table(doc, ["Metric", "Value"], rows)


def _format_redaction(doc: Document, data: dict):
    _format_anonymize(doc, data)


def _format_compare(doc: Document, data: dict):
    _add_title(doc, "Document Comparison Report")

    doc_a = data.get("document_a", "N/A")
    doc_b = data.get("document_b", "N/A")
    score = data.get("similarity_score")

    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, val) in enumerate([
        ("Document A", doc_a),
        ("Document B", doc_b),
        ("Similarity", f"{score}%" if score is not None else "N/A"),
    ]):
        cell = summary_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.size = Pt(10)
    doc.add_paragraph("")

    _add_label_value(doc, "Comparison Type", data.get("comparison_type", "General"))

    if data.get("summary"):
        _add_heading(doc, "Summary")
        _add_para(doc, data["summary"])

    sims = data.get("similarities", [])
    if sims:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Similarities")
        _add_bullet_list(doc, sims)

    diffs = data.get("differences", [])
    if diffs:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Differences")
        _add_bullet_list(doc, diffs)

    unique_a = data.get("unique_to_a", [])
    if unique_a:
        _add_horizontal_rule(doc)
        _add_heading(doc, f"Unique to {doc_a}")
        _add_bullet_list(doc, unique_a)

    unique_b = data.get("unique_to_b", [])
    if unique_b:
        _add_horizontal_rule(doc)
        _add_heading(doc, f"Unique to {doc_b}")
        _add_bullet_list(doc, unique_b)

    recs = data.get("recommendations", [])
    if recs:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Recommendations")
        _add_numbered_list(doc, recs)


def _format_merge(doc: Document, data: dict):
    _add_title(doc, "Merged Document")

    _add_label_value(doc, "Strategy", data.get("strategy", "N/A"))
    sources = data.get("source_documents", [])
    if sources:
        _add_label_value(doc, "Source Documents", ", ".join(str(s) for s in sources))

    outline = data.get("merged_outline", [])
    if outline:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Outline")
        _add_bullet_list(doc, outline)

    text = data.get("merged_text")
    if text:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Merged Content")
        for paragraph in str(text).split("\n\n"):
            paragraph = paragraph.strip()
            if paragraph:
                _add_para(doc, paragraph)

    notes = data.get("merge_notes")
    if notes:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Merge Notes")
        if isinstance(notes, list):
            _add_bullet_list(doc, notes)
        else:
            _add_para(doc, notes)


def _format_split(doc: Document, data: dict):
    _add_title(doc, f"Document Split: {data.get('document', 'Results')}")

    _add_label_value(doc, "Split Strategy", data.get("split_strategy", "N/A"))
    _add_label_value(doc, "Total Chunks", data.get("total_chunks", 0))

    chunks = data.get("chunks", [])
    if chunks:
        _add_horizontal_rule(doc)
        _add_heading(doc, "Chunks")
        for i, chunk in enumerate(chunks, 1):
            if isinstance(chunk, dict):
                title = chunk.get("title") or chunk.get("heading") or f"Chunk {i}"
                _add_heading(doc, title, level=2)
                text = chunk.get("text") or chunk.get("content", "")
                _add_para(doc, text)
                if chunk.get("word_count"):
                    p = doc.add_paragraph()
                    run = p.add_run(f"Word count: {chunk['word_count']}")
                    run.font.size = Pt(9)
                    run.font.color.rgb = MUTED_COLOR
            else:
                _add_heading(doc, f"Chunk {i}", level=2)
                _add_para(doc, str(chunk))


def _format_classify(doc: Document, data: dict):
    _add_title(doc, f"Classification: {data.get('document', 'Results')}")
    for k, v in data.items():
        if k == "document":
            continue
        if isinstance(v, list):
            _add_heading(doc, k.replace("_", " ").title())
            if v and isinstance(v[0], dict):
                headers = list(v[0].keys())
                rows = [[item.get(h, "") for h in headers] for item in v if isinstance(item, dict)]
                _add_table(doc, [h.replace("_", " ").title() for h in headers], rows)
            else:
                _add_bullet_list(doc, v)
        elif isinstance(v, dict):
            _add_heading(doc, k.replace("_", " ").title())
            for sk, sv in v.items():
                _add_label_value(doc, sk.replace("_", " ").title(), sv)
        else:
            _add_label_value(doc, k.replace("_", " ").title(), v)


def _format_custom(doc: Document, data: dict):
    _add_title(doc, "Processing Results")

    response = data.get("response")
    if response is not None:
        if isinstance(response, str):
            _add_heading(doc, "Response")
            for paragraph in response.split("\n\n"):
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                if paragraph.startswith("# "):
                    _add_heading(doc, paragraph[2:], level=1)
                elif paragraph.startswith("## "):
                    _add_heading(doc, paragraph[3:], level=2)
                elif paragraph.startswith("### "):
                    _add_heading(doc, paragraph[4:], level=3)
                elif paragraph.startswith("- ") or paragraph.startswith("* "):
                    items = [line.lstrip("- *").strip() for line in paragraph.split("\n") if line.strip()]
                    _add_bullet_list(doc, items)
                elif paragraph[0].isdigit() and ". " in paragraph[:5]:
                    items = []
                    for line in paragraph.split("\n"):
                        line = line.strip()
                        if line and line[0].isdigit() and ". " in line[:5]:
                            items.append(line.split(". ", 1)[1])
                        elif line:
                            items.append(line)
                    _add_numbered_list(doc, items)
                else:
                    _add_para(doc, paragraph)
        elif isinstance(response, dict):
            _add_heading(doc, "Response")
            _format_dict_recursive(doc, response, level=2)
        elif isinstance(response, list):
            _add_heading(doc, "Response")
            for i, item in enumerate(response, 1):
                if isinstance(item, dict):
                    _add_heading(doc, f"Item {i}", level=2)
                    for k, v in item.items():
                        _add_label_value(doc, k.replace("_", " ").title(), v)
                else:
                    _add_para(doc, f"{i}. {item}")
    else:
        _format_dict_recursive(doc, data, level=1)


def _format_dict_recursive(doc: Document, d: dict, level: int = 1):
    for k, v in d.items():
        if isinstance(v, dict):
            _add_heading(doc, k.replace("_", " ").title(), level=min(level, 4))
            _format_dict_recursive(doc, v, level + 1)
        elif isinstance(v, list):
            _add_heading(doc, k.replace("_", " ").title(), level=min(level, 4))
            if v and isinstance(v[0], dict):
                headers = list(v[0].keys())
                rows = [[item.get(h, "") for h in headers] for item in v if isinstance(item, dict)]
                _add_table(doc, [h.replace("_", " ").title() for h in headers], rows)
            else:
                str_items = [str(item) for item in v]
                if str_items:
                    _add_bullet_list(doc, str_items)
        else:
            _add_label_value(doc, k.replace("_", " ").title(), v)


def _append_original_content(doc: Document, original_text: str, filename: Optional[str] = None):
    doc.add_page_break()

    _add_title(doc, "Original Document Content")

    if filename:
        _add_label_value(doc, "Source File", filename)

    _add_horizontal_rule(doc)

    for paragraph in original_text.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        p = doc.add_paragraph()
        run = p.add_run(paragraph)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        p.paragraph_format.space_after = Pt(6)


TOOL_FORMATTERS = {
    "smart_summary": _format_summary,
    "smart_extract": _format_extract,
    "smart_translate": _format_translate,
    "smart_anonymize": _format_anonymize,
    "smart_audit": _format_audit,
    "smart_redaction": _format_redaction,
    "smart_compare": _format_compare,
    "smart_merge": _format_merge,
    "smart_split": _format_split,
    "smart_classify": _format_classify,
}


def format_result_to_docx(
    tool_name: str,
    data: dict,
    filename: Optional[str] = None,
    original_text: Optional[str] = None,
) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    _add_info_box(doc, tool_name, filename)

    formatter = TOOL_FORMATTERS.get(tool_name)
    if formatter:
        try:
            formatter(doc, data)
        except Exception as exc:
            logger.warning("Formatter for %s failed, falling back to generic: %s", tool_name, exc)
            _format_custom(doc, data)
    else:
        _format_custom(doc, data)

    if original_text and original_text.strip():
        _append_original_content(doc, original_text, filename)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
