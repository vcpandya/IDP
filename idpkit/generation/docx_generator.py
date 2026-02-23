"""DOCX document generator — converts content to Word format."""

import logging
import os
import re

logger = logging.getLogger(__name__)


def generate_docx(
    content: str,
    output_path: str,
    template_path: str = None,
) -> str:
    """Generate a DOCX file from markdown-formatted content.

    Parameters
    ----------
    content:
        Markdown-formatted text to convert.  Supports headings (``#``),
        bold (``**``), italic (``*``), bullet lists (``- ``), and
        simple markdown tables.
    output_path:
        Destination path for the generated ``.docx`` file.
    template_path:
        Optional path to a ``.docx`` template file.  If provided, the
        template is opened and content is appended to it.

    Returns
    -------
    str
        The absolute path to the generated file.

    Raises
    ------
    ImportError
        If python-docx is not installed.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX generation. "
            "Install it with: pip install python-docx"
        )

    logger.info("Generating DOCX: %s", output_path)

    if template_path and os.path.isfile(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # --- Headings ---
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            # python-docx heading levels are 0-based for Heading 1
            doc.add_heading(title, level=min(level, 9))
            i += 1
            continue

        # --- Markdown table ---
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, table_lines)
            continue

        # --- Bullet list ---
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, text)
            i += 1
            continue

        # --- Numbered list ---
        numbered_match = re.match(r"^\d+\.\s+(.+)$", line.strip())
        if numbered_match:
            text = numbered_match.group(1)
            para = doc.add_paragraph(style="List Number")
            _add_formatted_text(para, text)
            i += 1
            continue

        # --- Normal paragraph ---
        text = line.strip()
        if text:
            para = doc.add_paragraph()
            _add_formatted_text(para, text)
        i += 1

    # Ensure output directory exists.
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    doc.save(output_path)
    abs_path = os.path.abspath(output_path)
    logger.info("DOCX generated: %s", abs_path)
    return abs_path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_formatted_text(para, text: str) -> None:
    """Add text to a paragraph, handling bold and italic markers."""
    # Process **bold** and *italic* markers.
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)


def _add_markdown_table(doc, table_lines: list[str]) -> None:
    """Parse markdown table lines and add a Word table to the document."""
    rows_data: list[list[str]] = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Skip separator rows (e.g. | --- | --- |).
        if all(re.match(r"^:?-+:?$", c) for c in cells if c):
            continue
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows_data):
        for col_idx in range(num_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            table.rows[row_idx].cells[col_idx].text = cell_text
